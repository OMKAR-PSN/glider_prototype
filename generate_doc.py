"""
Generate Glider GNC Project Word Document
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Colour palette ────────────────────────────────────────────────────────
NAVY      = RGBColor(0x00, 0x27, 0x64)   # deep navy blue
ACCENT    = RGBColor(0x00, 0x6E, 0xC7)   # bright blue
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xF9)   # very light blue-grey


def set_run_font(run, size=11, bold=False, italic=False, colour=None):
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    if colour:
        run.font.color.rgb = colour


def shade_cell(cell, fill_hex: str):
    """Set cell background colour (hex string like '002764')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def add_cell_text(cell, text, bold=False, size=10, colour=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    set_run_font(run, size=size, bold=bold, colour=colour)


def styled_table(headers, rows, col_widths=None):
    """Insert a styled table with a navy header row."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, '002764')
        add_cell_text(cell, h, bold=True, size=10, colour=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = 'F0F4F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            shade_cell(cell, bg)
            add_cell_text(cell, cell_text, size=9.5)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return table


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, colour=NAVY)
    # Bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '002764')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, colour=ACCENT)
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, colour=DARK_GREY)
    return p


def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold, italic=italic, colour=DARK_GREY)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.3 + level * 0.3)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=10.5, colour=DARK_GREY)
    return p


def code_block(lines):
    """Monospaced grey box for code/diagrams."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # light grey shading
        rPr  = run._r.get_or_add_rPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'EFEFEF')
        rPr.append(shd)
    doc.add_paragraph()


def info_box(text):
    """Light-blue info callout."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell  = table.rows[0].cells[0]
    shade_cell(cell, 'E8F4FD')
    para  = cell.paragraphs[0]
    run   = para.add_run(text)
    set_run_font(run, size=10, italic=True, colour=RGBColor(0x00, 0x4E, 0x9A))
    doc.add_paragraph()


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AUTONOMOUS PARAFOIL GLIDER")
set_run_font(run, size=28, bold=True, colour=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GNC System — Complete Project Documentation")
set_run_font(run, size=16, bold=False, colour=ACCENT)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guidance  |  Navigation  |  Control  |  AI Training  |  Testing")
set_run_font(run, size=11, italic=True, colour=DARK_GREY)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Project GARUD  —  INSPACe 2026 Rocketry Competition")
set_run_font(run, size=11, bold=True, colour=DARK_GREY)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
heading1("Table of Contents")
toc_items = [
    ("1", "What Is This Project?"),
    ("2", "The Big Picture — Mission Overview"),
    ("3", "Hardware — What We Built With"),
    ("4", "Sensors — The Eyes and Ears"),
    ("5", "The State Machine — The Glider's Brain States"),
    ("6", "State Estimation — Knowing Where You Are"),
    ("7", "The Control Algorithm — How the Glider Steers"),
    ("8", "AI Training — Teaching the Glider to Fly"),
    ("9", "Terrain Mapping Integration"),
    ("10","4-Core Architecture — Running on Raspberry Pi 5"),
    ("11","Crash Recovery — What Happens If the Pi Restarts"),
    ("12","Testing — How We Verified Everything"),
    ("13","Key Results"),
    ("14","Glossary"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  {num}.  {title}")
    set_run_font(run, size=10.5, colour=DARK_GREY)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 1
# ═══════════════════════════════════════════════════════════════════════════
heading1("1.  What Is This Project?")
body("This is the Guidance, Navigation and Control (GNC) System for an autonomous parafoil glider. In simple words:")
info_box('We built the software "pilot" that controls a parachute-glider. When released from a rocket at high altitude, this system autonomously steers the glider to a pre-set landing target — without any human controlling it.')

heading2("Why a Parafoil and Not a Normal Parachute?")
body("A normal parachute just drifts wherever the wind takes it. A parafoil (also called a ram-air parachute) has a wing shape. By pulling the left or right brake line, you can steer it — just like steering a bicycle. This lets us guide it to a precise landing spot.")
code_block([
    "          +========================+",
    "          |   Parafoil Canopy     |",
    "          |   (inflated wing)     |",
    "          +========================+",
    "                  |         |",
    "          Left    |         |  Right",
    "          Brake   |         |  Brake",
    "          Line    |         |  Line",
    "                  |         |",
    "              +---+---------+---+",
    "              |      PAYLOAD    |   <- Raspberry Pi 5, sensors, servos",
    "              +-----------------+",
])

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 2
# ═══════════════════════════════════════════════════════════════════════════
heading1("2.  The Big Picture — Mission Overview")

heading2("Step-by-Step Mission Flow")
code_block([
    "  STEP 1        STEP 2        STEP 3         STEP 4         STEP 5",
    "  --------     --------      ----------     ----------     ----------",
    "  Rocket   ->  Reaches   ->  Ejects the ->  Parafoil   ->  Guided",
    "  launched     apogee        payload         deploys at     descent to",
    "               (top)         (glider)        600m AGL       landing target",
])

styled_table(
    ["Phase", "Altitude", "What Happens"],
    [
        ["Boost",                   "0 m -> Peak",    "Rocket motor burning. Glider is inside the rocket."],
        ["Drogue Descent",          "Peak -> 600m AGL","Small drogue parachute slows the payload. Glider waiting."],
        ["Deployment Trigger",      "600m AGL",        "Servo fires, releasing the main parafoil canopy."],
        ["Deployment Verification", "~600m",           "Software confirms canopy is fully open."],
        ["Guided Descent",          "600m -> 0m",      "GNC algorithm takes over and steers to the target."],
        ["Landed",                  "< 5m AGL",        "Both brakes pulled (flare). Mission complete."],
    ],
    col_widths=[1.5, 1.3, 3.8]
)

body("AGL = Above Ground Level (height above the launch site, not sea level).", italic=True)

heading2("Key Safety Feature — Deployment Moving Average")
body("The 600m trigger does NOT fire on a single sensor reading. It requires all three conditions:")
bullet("10 barometer readings averaged together (filters out noise/spikes)")
bullet("The averaged altitude must be <= 600m AGL")
bullet("Vertical velocity must confirm descent (< -2 m/s)")
body("This prevents accidental early deployment from a momentary bad sensor reading.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 3
# ═══════════════════════════════════════════════════════════════════════════
heading1("3.  Hardware — What We Built With")

heading2("Main Computer")
styled_table(
    ["Item", "Detail"],
    [
        ["Computer", "Raspberry Pi 5 (4 CPU cores — each used for a different task)"],
        ["Power Monitor", "INA219 — monitors battery voltage and current"],
    ],
    col_widths=[1.5, 5.0]
)

heading2("Actuators (Things That Move)")
body("These are the 'muscles' of the system — they physically pull the brake lines of the parafoil.")
styled_table(
    ["Component", "Model", "PCA9685 Channel", "What It Does"],
    [
        ["Left Brake Servo",    "EMAX ES3004", "Ch 0", "Pulls left brake line to turn left"],
        ["Right Brake Servo",   "EMAX ES3004", "Ch 1", "Pulls right brake line to turn right"],
        ["Drogue Release",      "EMAX ES3004", "Ch 2", "Releases the main parafoil at 600m AGL"],
        ["Gimbal Roll Servo",   "28BYJ-48",    "Ch 3", "Stabilises camera/payload gimbal (roll axis)"],
        ["Gimbal Pitch Servo",  "MG90",        "Ch 4", "Stabilises camera/payload gimbal (pitch axis)"],
    ],
    col_widths=[1.6, 1.3, 1.3, 2.9]
)

info_box("PCA9685 = A chip that lets you control up to 16 servos using just 2 wires (I2C bus). All servos are controlled through this chip, not directly from the Pi GPIO pins.")

heading2("Servo Angle Convention")
body("Servos rotate between 60 degrees and 120 degrees. 90 degrees is the center/neutral position.")
code_block([
    "       60 deg           90 deg           120 deg",
    "       |----------------|----------------|",
    "    Full Left          Neutral         Full Right",
    "   (hard turn)       (straight)        (hard turn)",
])

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 4
# ═══════════════════════════════════════════════════════════════════════════
heading1("4.  Sensors — The Eyes and Ears")
body("The glider needs to know three things at all times: WHERE is it (GPS), WHICH WAY is it pointing (IMU), and HOW HIGH is it (Barometer).")

styled_table(
    ["Sensor", "Model", "Interface", "What It Measures"],
    [
        ["IMU",                "BNO085",  "I2C 0x4A",        "Roll, Pitch, Yaw (orientation); Gyroscope; Accelerometer; Magnetometer"],
        ["Barometer (Primary)","BMP388",  "I2C 0x76",        "Air pressure -> Altitude (primary)"],
        ["Barometer (Backup)", "BMP388",  "I2C 0x77",        "Redundant altitude measurement (median vote)"],
        ["GPS",                "NEO-M8N", "UART /dev/ttyAMA0","Latitude, Longitude, Speed, Heading"],
        ["Power Monitor",      "INA219",  "I2C 0x41",        "Battery voltage and current draw"],
    ],
    col_widths=[1.5, 1.1, 1.5, 3.5]
)

heading2("Why Two Barometers?")
body("Single-point-of-failure is dangerous in flight. If one barometer gives a bad reading (e.g., a pressure spike), the software takes the median (middle value) of both. This gives us a reliable reading even if one sensor is faulty.")

heading2("What is an IMU?")
body("IMU = Inertial Measurement Unit. It is a single chip containing three sensors:")
bullet("Accelerometer — measures acceleration forces (like a phone tilt sensor)")
bullet("Gyroscope — measures how fast the device is rotating")
bullet("Magnetometer — a digital compass that gives absolute heading")
body("The BNO085 is special because it has a built-in fusion processor that directly outputs ready-to-use roll, pitch, and yaw angles — no extra calculations needed on the Pi.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 5
# ═══════════════════════════════════════════════════════════════════════════
heading1("5.  The State Machine — The Glider's Brain States")
body("The state machine is like a strict traffic light. The glider can only be in ONE state at a time, and it follows strict rules to move between them.")

code_block([
    "  [BOOST]",
    "      |  altitude drops + velocity < -2 m/s",
    "      v",
    "  [DROGUE_DESCENT]",
    "      |  average altitude <= 600m AGL AND velocity < -2 m/s",
    "      v",
    "  [DEPLOYMENT_TRIGGER]  --> fires drogue servo",
    "      |  immediately",
    "      v",
    "  [DEPLOYMENT_VERIFICATION]  --> confirms canopy open",
    "      |  immediately",
    "      v",
    "  [GUIDED_DESCENT]  <-- GNC algorithm ACTIVE here",
    "      |  altitude <= 5m AGL",
    "      v",
    "  [LANDED]",
])

styled_table(
    ["State", "What the System Is Doing"],
    [
        ["BOOST",                   "Inside rocket, no control needed. Waiting for apogee."],
        ["DROGUE_DESCENT",          "Falling under small drogue chute. Waiting for 600m AGL."],
        ["DEPLOYMENT_TRIGGER",      "600m reached. Fires the servo that releases the parafoil."],
        ["DEPLOYMENT_VERIFICATION", "Waits one cycle to confirm canopy is inflated."],
        ["GUIDED_DESCENT",          "Full autonomous steering to landing target (AI / PID active)."],
        ["LANDED",                  "< 5m AGL detected. Mission complete. Log and stop."],
    ],
    col_widths=[2.0, 5.5]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 6
# ═══════════════════════════════════════════════════════════════════════════
heading1("6.  State Estimation — Knowing Where You Are")
body("Raw sensor data is noisy — a barometer reading might randomly jump by 2 metres. We use mathematical filters to get smooth, accurate estimates of position and orientation.")

heading2("6.1  Attitude Estimation — Madgwick Filter (SITL) / BNO085 (Real Hardware)")
body("In simulation mode, the Madgwick filter combines accelerometer + gyroscope + magnetometer data to compute roll, pitch, and yaw. On real hardware, the BNO085's built-in AHRS processor does this directly.")
info_box("Gyroscope: fast but drifts over time (like a spinning top that slowly wobbles). Accelerometer: slow but accurate long-term. Madgwick filter blends both for a result that is fast AND accurate.")

heading2("6.2  Altitude Estimation — Extended Kalman Filter (EKF)")
body("The EKF combines the barometer (accurate but noisy) and the accelerometer (fast but drifts) to produce a smooth altitude and vertical velocity estimate. It runs every loop cycle (20 Hz):")
bullet("Predict step: Uses accelerometer to estimate new altitude")
bullet("Update step: Corrects the prediction using the barometer reading")
body("The EKF also gives us vertical velocity (negative = descending), which is critical for the 600m deployment trigger check.")

heading2("6.3  Wind Estimation — Recursive Least Squares (RLS)")
body("Wind pushes the glider sideways. Knowing the wind lets the algorithm fight it. The RLS estimator continuously learns the wind vector from the difference between:")
bullet("Where the glider's nose is pointing (from IMU heading)")
bullet("Where it is actually moving over the ground (from GPS velocity)")
body("If GPS says you are moving east but your nose points north — wind is blowing you east. The estimator captures this and feeds it to the guidance algorithm.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 7
# ═══════════════════════════════════════════════════════════════════════════
heading1("7.  The Control Algorithm — How the Glider Steers")
body("There are TWO controllers. The AI runs first; the traditional PID acts as an automatic backup if anything goes wrong.")

heading2("7.1  The Heading Error — Knowing How Much to Turn")
body("First, the software calculates how much the glider needs to turn:")
code_block([
    "  heading_error = target_bearing - current_heading",
    "",
    "  If heading_error > 0  --> need to turn RIGHT",
    "  If heading_error < 0  --> need to turn LEFT",
    "  If heading_error = 0  --> pointing directly at target",
])
body("The error is always kept in the range [-180 deg, +180 deg] so the glider never turns the long way around.")

heading2("7.2  Backup Controller — Heading PID")
body("The PID controller outputs a brake deflection command based on the heading error:")
code_block([
    "  delta_a = Kp * heading_error",
    "           + Ki * (integral of heading_error over time)",
    "           + Kd * (rate of change of heading_error)",
])
body("Where Kp, Ki, Kd are tuning gains loaded from config/gains.yaml. These gains automatically change based on altitude (gain scheduling):")
styled_table(
    ["Altitude Phase", "Gains", "Reason"],
    [
        ["Cruise (high altitude)",  "Gentle",     "Plenty of time and altitude — no rush"],
        ["Approach (medium)",       "Medium",     "Start tightening the turn"],
        ["Final (< 50m AGL)",       "Aggressive", "Must reach target quickly before ground impact"],
    ],
    col_widths=[1.8, 1.0, 3.8]
)

heading2("7.3  Adaptive Gain Control (AGC) — Fighting the Wind")
body("The gain Kp is not fixed — it adapts based on crosswind strength:")
code_block([
    "  K_wind = BASE_GAIN + (K_MAX - K_MIN) * |crosswind| / (|crosswind| + sigma)",
    "",
    "  BASE_GAIN = 15.0   (gain in calm air)",
    "  K_MAX     = 25.0   (maximum gain in strong crosswind)",
    "  K_MIN     =  5.0   (minimum gain floor)",
    "  sigma     = 30.0   (sensitivity shaping constant)",
])
body("In calm air: gain = 15.0 (gentle corrections). In strong crosswind: gain increases to 25.0 (fights harder to stay on course). This prevents the glider from being blown away without overreacting in calm conditions.")

heading2("7.4  The Elevon Mixer — Brake Commands to Servo Angles")
body("The PID gives delta_a (differential brake) and delta_s (symmetric brake / speed control). The mixer converts these to physical servo angles:")
code_block([
    "  Left_Servo  = 90 + delta_s - delta_a    (clamped to [60, 120] deg)",
    "  Right_Servo = 90 + delta_s + delta_a    (clamped to [60, 120] deg)",
])
body("Worked Example — Target is 30 degrees to the right:")
code_block([
    "  heading_error = +30 deg  (need to turn right)",
    "  PID outputs:  delta_a = +12 deg,  delta_s = 0 deg",
    "  Left servo:   90 + 0 - 12 = 78 deg  (slightly relaxed)",
    "  Right servo:  90 + 0 + 12 = 102 deg (slightly pulled in)",
    "  Result:       right brake pulled more --> glider turns right  [CORRECT]",
])

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 8
# ═══════════════════════════════════════════════════════════════════════════
heading1("8.  AI Training — Teaching the Glider to Fly")

heading2("8.1  What is Reinforcement Learning?")
body("Instead of hand-writing every steering rule, we trained an AI agent to learn how to fly the glider by itself — the same way you would train a dog using treats.")
info_box("Dog fetches ball --> give it a treat (reward). Dog runs away --> no treat. Over time, dog learns fetching = treats. We do the same thing: glider moves toward target --> reward. Glider misses --> penalty. After millions of practice runs in simulation, it learns to land precisely.")

heading2("8.2  The Algorithm — SAC (Soft Actor-Critic)")
body("We chose SAC because it is designed for smooth, continuous control — not just ON/OFF decisions. It learns to output precise servo angle commands.")

heading2("8.3  What the AI Sees — 16 Observation Inputs")
styled_table(
    ["Index", "Signal", "Plain English Meaning"],
    [
        ["0-1",  "sin/cos of heading error",     "How much to turn and in which direction"],
        ["2",    "Distance to target",            "How far away is the landing zone"],
        ["3",    "Altitude excess",               "Do we have too much or too little altitude to reach the target"],
        ["4",    "Wind speed",                    "How strong is the wind blowing"],
        ["5-6",  "sin/cos of wind direction",     "Which direction the wind is blowing from"],
        ["7",    "Pitch rate",                    "Is the nose pitching up or down"],
        ["8",    "Roll rate",                     "Is the glider rolling left or right"],
        ["9",    "Yaw rate",                      "Is the glider spinning"],
        ["10-11","Previous servo commands",        "Smoothness signal — avoids sudden jerky movements"],
        ["12-13","sin/cos of track error",         "Difference between where nose points and where glider travels (wind crab angle)"],
        ["14",   "Lateral drift rate",             "How fast wind is pushing the glider sideways"],
        ["15",   "Time to impact",                 "How many seconds until we hit the ground"],
    ],
    col_widths=[0.6, 1.9, 4.0]
)
body("sin/cos encoding is used for angles to avoid the wrap-around problem where 359 deg and 0 deg look very different as numbers but are actually the same direction.", italic=True)

heading2("8.4  What the AI Outputs — 2 Control Commands")
styled_table(
    ["Output", "Range", "Meaning"],
    [
        ["delta_a (asymmetric brake)", "-30 deg to +30 deg", "Left/right differential steering"],
        ["delta_s (symmetric brake)",  "0 deg to +30 deg",   "Slows down the glider (both brakes together)"],
    ],
    col_widths=[2.0, 1.5, 3.0]
)

heading2("8.5  Curriculum Learning — Teaching Step by Step")
body("We didn't start training with the hardest scenario immediately. Like a student progressing from easy to hard problems:")
styled_table(
    ["Stage", "Conditions", "Success Criterion"],
    [
        ["Stage 1", "Calm wind, target close by",           "Land within 50m"],
        ["Stage 2", "Some wind, more distance to target",   "Land within 50m"],
        ["Stage 3", "Strong wind, far target (full difficulty)", "Land within 50m"],
    ],
    col_widths=[0.8, 2.7, 2.0]
)

heading2("8.6  The Reward Function")
body("The AI gets a score every 0.05 seconds (every loop cycle):")
code_block([
    "  Reward = (previous distance - current distance)    <- reward for getting closer",
    "          - 0.05 * |change in servo commands|        <- penalty for jerky movement",
    "",
    "  Bonus on landing:",
    "    Within 10m  --> +10,000 points",
    "    Within 20m  --> +2,000 points",
    "    Within 50m  --> +200 points",
])

heading2("8.7  Training Results")
styled_table(
    ["Checkpoint", "Training Steps", "CEP50", "Success Rate"],
    [
        ["7M checkpoint (12D obs — archived)", "7,000,000", "333m",  "2.3%"],
        ["6.5M checkpoint (16D obs — ACTIVE)", "6,500,000", "106m",  "6.1%"],
    ],
    col_widths=[2.5, 1.4, 0.8, 1.3]
)
body("CEP50 = Circular Error Probable at 50% — the radius of the circle that contains 50% of all landings.", italic=True)
body("The jump from 333m to 106m happened when we added 4 new wind-aware signals (track error, lateral drift) to the observation space. The AI could finally compute a wind-corrected course.")

heading2("8.8  Deploying the AI on Raspberry Pi")
body("Training happens on a PC using PyTorch. For the Raspberry Pi, we export the trained model to ONNX format — a lightweight, portable format that any computer can run.")
code_block([
    "  PC Training (PyTorch) --> Export to .onnx --> Raspberry Pi (ONNX Runtime)",
    "                                                Inference: 0.019ms per step",
    "                                                Budget:    5ms per step",
])
info_box("Safety: If the ONNX model fails for ANY reason (crash, NaN output, slow response, GPS stale, shape mismatch), the PID controller automatically and silently takes over within the same loop cycle. Zero gap in control.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 9
# ═══════════════════════════════════════════════════════════════════════════
heading1("9.  Terrain Mapping Integration")
body("The GARUD terrain mapping system runs on a separate CPU core from the GNC system. It continuously processes terrain data to ensure the glider avoids obstacles and selects the safest flight path.")
bullet("3D terrain data of the landing area (hills, trees, buildings)")
bullet("Obstacle detection and avoidance waypoint generation")
bullet("Optimal path suggestion to the safest landing zone within range")
body("The GNC and terrain mapping systems communicate through a shared inter-process memory queue (Python multiprocessing Queue). If the primary landing zone is obstructed by terrain, the terrain mapper provides an alternative waypoint to the GNC system.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 10
# ═══════════════════════════════════════════════════════════════════════════
heading1("10.  4-Core Architecture — Raspberry Pi 5")
body("The Raspberry Pi 5 has 4 CPU cores. Instead of running everything in one slow program, we split the work between cores to ensure nothing blocks anything else.")

code_block([
    "  +----------------------------------------------------------------+",
    "  |                    Raspberry Pi 5                              |",
    "  |                                                                |",
    "  |  Core 0: SENSORS       Core 1: GNC          Core 2: TERRAIN  |",
    "  |  +--------------+      +-------------+      +------------+    |",
    "  |  | Read IMU     |      | State Mach. |      | Terrain    |    |",
    "  |  | Read GPS     |--+-->| AI / PID    |      | Map Proc.  |    |",
    "  |  | Read Baro    |  |   | Servo Output|      |            |    |",
    "  |  | @ 50 Hz      |  |   | @ 20 Hz     |      |            |    |",
    "  |  +--------------+  |   +-------------+      +------------+    |",
    "  |                    |                                           |",
    "  |  Core 3: LOGGING   |                                           |",
    "  |  +--------------+  |                                           |",
    "  |  | Write CSV    |<-+                                           |",
    "  |  | Write State  |                                              |",
    "  |  | SD Card I/O  |                                              |",
    "  |  +--------------+                                              |",
    "  +----------------------------------------------------------------+",
])

styled_table(
    ["Core", "Process", "Frequency", "Job"],
    [
        ["Core 0", "Sensor Process",  "50 Hz",    "Reads all sensors, puts data in shared memory"],
        ["Core 1", "GNC Process",     "20 Hz",    "Reads sensor data, runs AI/PID, writes servo angles"],
        ["Core 2", "Terrain Process", "As needed","Processes terrain maps, detects obstacles, refines waypoints"],
        ["Core 3", "Logger Process",  "10 Hz",    "Reads all data from queue, writes to SD card"],
    ],
    col_widths=[0.7, 1.4, 1.1, 4.4]
)
body("Each process is pinned to its specific core using Linux CPU affinity (os.sched_setaffinity). This prevents the operating system from moving processes between cores, ensuring predictable real-time timing.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 11
# ═══════════════════════════════════════════════════════════════════════════
heading1("11.  Crash Recovery — What Happens If the Pi Restarts")
body("A mid-flight power glitch or software crash cannot be allowed to cause double-deployment of the parachute or loss of control. We built a robust recovery system.")

heading2("How It Works")
body("Every 5 seconds during flight, the software writes a .state file to the SD card containing:")
bullet("Current flight phase (e.g., GUIDED_DESCENT)")
bullet("Whether the drogue has already been fired (drogue_fired = True/False)")
bullet("Last known altitude and vertical velocity")
bullet("Target coordinates")

code_block([
    "  Normal cold boot (no .state file):",
    "  [BOOT] --> Start fresh from BOOST state",
    "",
    "  After a Pi crash and restart (.state file found):",
    "  [BOOT] --> Read .state file",
    "          --> Jump directly to GUIDED_DESCENT",
    "          --> LOCK OUT drogue channel (already fired, don't fire again!)",
    "          --> Seed EKF with last known altitude (smooth startup)",
    "          --> Continue mission as if nothing happened",
])

info_box("Critical Safety: Once the drogue fires, drogue_fired = True is written to the .state file. On any restart, this flag is read and the drogue channel is PERMANENTLY locked out. This prevents the drogue from firing a second time even if the Pi restarts multiple times.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 12
# ═══════════════════════════════════════════════════════════════════════════
heading1("12.  Testing — How We Verified Everything")

heading2("12.1  Software-In-The-Loop (SITL)")
body("Before touching any real hardware, we tested everything in simulation. A 4-DOF physics engine simulates glider aerodynamics, wind, and all sensor noise.")
styled_table(
    ["SITL Result", "Value"],
    [
        ["Simulation duration",  "191 seconds"],
        ["State transitions",    "DROGUE -> DEPLOYMENT -> GUIDED -> LANDED (all correct)"],
        ["AI controller steps",  "3,713 (100% of guided descent)"],
        ["PID fallback steps",   "0 (AI ran cleanly throughout)"],
        ["Miss distance",        "589m (3 m/s crosswind from 1000m starting offset — within expected range)"],
    ],
    col_widths=[2.0, 5.5]
)

heading2("12.2  Monte Carlo Analysis (500 Simulated Drops)")
body("We ran 500 simulated drops at random starting positions and wind conditions to get statistical performance data.")
styled_table(
    ["Metric", "Value", "Meaning"],
    [
        ["CEP50", "106m", "50% of all landings fell within 106m of the target"],
        ["CEP90", "342m", "90% of all landings fell within 342m of the target"],
        ["Success Rate", "6.1%", "6.1% of landings fell within 50m of the target"],
    ],
    col_widths=[1.2, 0.8, 5.0]
)

heading2("12.3  Real Data Actuation Test — UAV-SEAD Dataset")
body("We tested the steering algorithm against real drone flight data downloaded from HuggingFace (PX4 ULog format).")
bullet("Dataset: UAV-SEAD (10 real PX4 flight log files, 2018-06-04)")
bullet("Total timesteps tested: 9,019 real sensor readings")
bullet("Real data used per timestep: actual yaw/heading, real barometric altitude, real wind")

styled_table(
    ["Check", "What Was Verified", "Result"],
    [
        ["Servo Limits",   "All servo angles stayed within [60 deg, 120 deg]",                  "ALL PASS"],
        ["Control Limits", "All delta_a values stayed within [-30 deg, +30 deg]",               "ALL PASS"],
        ["No NaN / Inf",   "Zero numerical errors across all 9,019 timesteps",                  "ALL PASS"],
        ["Sign Correct",   "Right heading error always produced a right turn (and vice versa)",  "ALL PASS"],
    ],
    col_widths=[1.3, 4.2, 1.2]
)

heading2("12.4  Unit Tests")
styled_table(
    ["Test File", "What It Verified"],
    [
        ["test_gnc.py",                  "Math conversions, bearing calculations, wrap-angle functions"],
        ["test_curriculum.py",           "Curriculum stage advancement logic"],
        ["test_buffer_flush.py",         "AI replay buffer correctly keeps newest experiences during stage transitions"],
        ["test_bmp388_hw.py",            "Physical barometer reads valid altitude on real hardware"],
        ["test_bno085_hw.py",            "Physical IMU gives valid roll/pitch/yaw on real hardware"],
        ["test_actuation_uav_dataset.py","Full AGC algorithm tested against real PX4 flight data"],
    ],
    col_widths=[2.2, 5.3]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 13
# ═══════════════════════════════════════════════════════════════════════════
heading1("13.  Key Results")

styled_table(
    ["Metric", "Value"],
    [
        ["Best AI checkpoint",           "6.5 million training steps (16D observation space)"],
        ["Landing accuracy (CEP50)",     "106 metres"],
        ["Landing accuracy (CEP90)",     "342 metres"],
        ["AI inference time (mean)",     "0.019ms per loop cycle (budget: 5ms)"],
        ["Actuation test samples",       "9,019 real sensor timesteps from 10 PX4 flight logs"],
        ["All actuation safety checks",  "PASS (servo limits, control limits, no NaN, sign correct)"],
        ["SITL end-to-end test",         "PASS (DROGUE -> DEPLOYMENT -> GUIDED -> LANDED, 0 PID fallbacks)"],
        ["ONNX model verification",      "PASS (shape, determinism, latency, fallback injection)"],
    ],
    col_widths=[2.5, 5.0]
)

heading2("Pre-Flight Confidence Checklist")
bullet("[PASS] AI model shape verified: input [1, 16], output [1, 2]")
bullet("[PASS] AI output bit-identical across 5 independent runs (deterministic)")
bullet("[PASS] PID fallback fires correctly on GPS stale, NaN output, inference timeout")
bullet("[PASS] Servo limits enforced in both hardware and software layers")
bullet("[PASS] Drogue double-fire prevented by crash recovery system")
bullet("[PASS] All state transitions verified in SITL end-to-end")
bullet("[PASS] ONNX vs PyTorch cross-check: max difference < 1e-04 (within tolerance)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 14 — GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════
heading1("14.  Glossary")
styled_table(
    ["Term", "Full Form / Meaning"],
    [
        ["AGL",       "Above Ground Level — height above the launch site (not sea level)"],
        ["AGC",       "Adaptive Gain Control — the method of adjusting PID gains based on wind/altitude"],
        ["AHRS",      "Attitude and Heading Reference System — onboard fusion processor in BNO085"],
        ["CEP50",     "Circular Error Probable at 50% — radius containing 50% of landings"],
        ["COG",       "Course Over Ground — actual direction of travel over the ground (may differ from heading due to wind)"],
        ["EKF",       "Extended Kalman Filter — mathematical sensor fusion for smooth altitude/velocity"],
        ["GNC",       "Guidance, Navigation, Control — the three main functions of a flight computer"],
        ["IMU",       "Inertial Measurement Unit — accelerometer + gyroscope + magnetometer combined"],
        ["NED",       "North-East-Down — the standard coordinate frame used in aviation"],
        ["ONNX",      "Open Neural Network Exchange — portable AI model format for deployment"],
        ["PID",       "Proportional-Integral-Derivative — the classic feedback control algorithm"],
        ["RLS",       "Recursive Least Squares — the math used for online wind estimation"],
        ["SAC",       "Soft Actor-Critic — the reinforcement learning algorithm used for AI training"],
        ["SITL",      "Software-In-The-Loop — running real flight software in a simulation environment"],
        ["delta_a",   "Asymmetric brake deflection — differential left/right braking for steering"],
        ["delta_s",   "Symmetric brake deflection — both brakes together for speed control"],
        ["psi",       "Heading angle — the compass direction the glider's nose points"],
        ["Delta psi", "Heading error — the angle between current heading and target bearing"],
    ],
    col_widths=[1.2, 6.3]
)

# ─── Save ──────────────────────────────────────────────────────────────────
OUT = r"C:\Users\Omkar\Desktop\Rocket\glider\Glider_GNC_Project_Documentation.docx"
doc.save(OUT)
print(f"Document saved to:\n{OUT}")
